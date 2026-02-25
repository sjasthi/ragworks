# This project implements a RAG framework that allows users to:
<<<<<<< HEAD
# 1. Upload a set of documents (PDFs, Word, HTML, text, etc). Dont use quotes with docx
# 2. Automatically preprocess and embed them into a retrievable knowledge store
# 3. Ask questions in natural language and get grounded, source-referenced responses


=======
# 1. Upload a set of documents (PDFs, Word, HTML, text, etc)
# 2. Automatically preprocess and embed them into a retrievable knowledge store
# 3. Ask questions in natural language and get grounded, source-referenced responses

>>>>>>> 84ff12a2e0009c059180310591bd9e2e38b2253c
import os
from openai import OpenAI, AuthenticationError
from dotenv import load_dotenv
from bs4 import BeautifulSoup
from docx import Document
import chromadb
from chromadb.utils import embedding_functions
<<<<<<< HEAD

print("RUNNING FILE:", __file__)
=======
from chromadb.config import Settings
>>>>>>> 84ff12a2e0009c059180310591bd9e2e38b2253c

# OpenAI Client
def create_client(api_key):
    try:
        client = OpenAI(api_key=api_key)
        client.models.list()
        return client
    except AuthenticationError:
<<<<<<< HEAD
        print("Invalid API key.")
        return None

=======
        print("Invalid API")
    return None
>>>>>>> 84ff12a2e0009c059180310591bd9e2e38b2253c
load_dotenv()

# Load api key
api_key = os.environ.get("OPENAI_KEY")
if not api_key:
    api_key = input("Enter your OpenAI API key: ").strip()

client = create_client(api_key)

<<<<<<< HEAD

# Persistent Chroma Setup
persist_dir = os.path.abspath("chroma_storage")
print("PERSIST DIR:", persist_dir)

chroma_client = chromadb.PersistentClient(path=persist_dir)

embedding_function = embedding_functions.OpenAIEmbeddingFunction(
=======
# Persistent ChromaDB only while the script is running
script_dir = os.path.dirname(os.path.abspath(__file__))
persist_dir = os.path.join(script_dir, "chroma_storage")

chroma_client = chromadb.Client(
    Settings(persist_directory=persist_dir, anonymized_telemetry=False)
)

openai_ef = embedding_functions.OpenAIEmbeddingFunction(
>>>>>>> 84ff12a2e0009c059180310591bd9e2e38b2253c
    api_key=api_key,
    model_name="text-embedding-3-small"
)

database = chroma_client.get_or_create_collection(
    name="RAG_Work_test",
<<<<<<< HEAD
    embedding_function=embedding_function
)

print("Current stored chunks:", database.count())


### Helper Functions

# Normalize file paths to ensure consistent storage and retrieval
def normalize_path(path):
    return os.path.abspath(path)


# Duplicate helper function to check if a document with the same source already exists in the database
def file_already_exists(file_path):
    file_path = normalize_path(file_path)
    results = database.get(where={"source": file_path}, limit=1)
    return len(results["ids"]) > 0

# Simple text chunking function with overlap to maintain context across chunks
def chunk_text(text, chunk_size=800, overlap=150):
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunks.append(text[start:end])
        start += chunk_size - overlap
    return chunks


### Ingestion

# Each add function checks if the file has already been uploaded (based on normalized path) to prevent duplicates.
# If not, it processes the file, chunks the text, and adds it to the database with metadata referencing the source file.

#html ingestion - using soup to extract the text from it

def add_html(file_path):
    file_path = normalize_path(file_path)

    if file_already_exists(file_path):
        print("File already uploaded.")
        return

=======
    embedding_function=openai_ef
)

print("Persistent Chroma collection at:", persist_dir)
print("Total stored chunks so far:", database.count())

# HTML ingestion - using soup to extract the text from it
def add_html(file_path):
>>>>>>> 84ff12a2e0009c059180310591bd9e2e38b2253c
    with open(file_path, "r", encoding="utf-8") as f:
        content = f.read()

    soup = BeautifulSoup(content, "html.parser")
<<<<<<< HEAD
    text = soup.get_text(separator=" ", strip=True)

    chunks = chunk_text(text)
    file_name = os.path.splitext(os.path.basename(file_path))[0]

    database.add(
        documents=chunks,
        ids=[f"{file_name}_chunk_{i}" for i in range(len(chunks))],
        metadatas=[{"source": file_path} for _ in chunks]
    )

    print("Added successfully. Collection size:", database.count())

#docx ingestion
def add_docx(file_path):
    file_path = normalize_path(file_path)

    if file_already_exists(file_path):
        print("File already uploaded.")
        return

    doc = Document(file_path)
    text = "\n".join([p.text for p in doc.paragraphs])

    chunks = chunk_text(text)
    file_name = os.path.splitext(os.path.basename(file_path))[0]

    database.add(
        documents=chunks,
        ids=[f"{file_name}_chunk_{i}" for i in range(len(chunks))],
        metadatas=[{"source": file_path} for _ in chunks]
    )

    print("Added successfully. Collection size:", database.count())

#txt Ingestion
def add_txt(file_path):
    file_path = normalize_path(file_path)

    if file_already_exists(file_path):
        print("File already uploaded.")
        return

    with open(file_path, "r", encoding="utf-8") as f:
        text = f.read()

    chunks = chunk_text(text)
    file_name = os.path.splitext(os.path.basename(file_path))[0]

    database.add(
        documents=chunks,
        ids=[f"{file_name}_chunk_{i}" for i in range(len(chunks))],
        metadatas=[{"source": file_path} for _ in chunks]
    )

    print("Added successfully. Collection size:", database.count())


# Retrieval
def get_context(query, k=3):
    results = database.query(
        query_texts=[query],
        n_results=k  # k gives top 3 relevant matches to the query
    )
    # querying the database with relevant info then extracting
    documents = results["documents"][0]
    metadatas = results["metadatas"][0]

    context_blocks = []
    for doc, meta in zip(documents, metadatas):
        context_blocks.append(f"[Source: {meta['source']}]\n{doc}")

    return "\n\n".join(context_blocks)

=======
    htm_txt = soup.get_text(strip=True)

    database.add(
        documents=[htm_txt],
        metadatas=[{"source": file_path}],
        ids=[f"html_{os.path.basename(file_path)}"]
    )

    print(f"{file_path} added successfully. Total stored chunks now: {database.count()}")

# DOCX ingestion
def add_docx(file_path):
    file_name = os.path.splitext(os.path.basename(file_path))[0]

    # Load DOCX
    doc = Document(file_path)
    text = "\n".join([p.text for p in doc.paragraphs])

    # Chunk Text
    chunk_size = 500
    chunks = [text[i:i+chunk_size] for i in range(0, len(text), chunk_size)]

    # Create Embeddings
    response = client.embeddings.create(
        model="text-embedding-3-small",
        input=chunks
    )
    embeddings = [item.embedding for item in response.data]

    # Add to collection with unique IDS and metadata
    database.add(
        documents=chunks,
        embeddings=embeddings,
        ids=[f"{file_name}_chunk_{i}" for i in range(len(chunks))],
        metadatas=[{"source": file_name} for _ in chunks]
    )

    print(f"{file_path} added successfully. Total stored chunks now: {database.count()}")

# TXT Ingestion
def add_txt(file_path):
    with open(file_path, "r", encoding="utf-8") as file:
        content = file.read()

    chunks = [sentence.strip() for sentence in content.split(".") if sentence]

    print("Chunks:", chunks)

    database.add(
        documents=chunks,
        ids=[f"txt_doc{i}" for i in range(len(chunks))],
        metadatas=[{"source": file_path}] * len(chunks)
    )

    print(f"{file_path} added successfully. Total stored chunks now: {database.count()}")

# Retrieval
def get_context(query):
    results = database.query(
        query_texts=[query],
        n_results=3, # gives top 3 relevant matches to the query
    )
    # querying the database with relevant info then extracting
    if results["documents"] and results["documents"][0]:
        context = "\n\n".join(results["documents"][0])
        return context
    return ""
>>>>>>> 84ff12a2e0009c059180310591bd9e2e38b2253c

# Chat
def chat(user_input):
    context = get_context(user_input)

<<<<<<< HEAD
    messages = [
        {
            "role": "system",
            "content": "Answer the question using ONLY the provided context. If not in context, say you do not know."
        },
        {
            "role": "user",
            "content": f"Context:\n{context}\n\nQuestion: {user_input}"
        }
    ]

    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=messages,
    )

    return response.choices[0].message.content


# Loop for uploading files
while True:
    file_type = input("\nWhat type of file do you want to add? (html/docx/txt/none): ").strip().lower()
    if file_type == "none":
        break

    file_path = input("Enter file path: ").strip().strip('"')  # Remove quotes if user included them
=======
    msg = [{
        "role": "developer",
        "content": "Talk like an assistant, use the context to answer the question."
    }]

    if context:
        msg.append({
            "role": "user",
            "content": f"Context:\n{context}\n\nQuestion: {user_input}"
        })
    else:
        msg.append({"role": "user", "content": user_input})

    comp = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=msg,
    )
    response = comp.choices[0].message.content
    return response

# Loop for uploading files
while True:
    add_file = input("\nWhat type of file do you want to add? (html/docx/txt/none): ").strip().lower()
    if add_file == "none":
        break

    file_path = input("Enter the path to your file: ").strip()
>>>>>>> 84ff12a2e0009c059180310591bd9e2e38b2253c

    if not os.path.exists(file_path):
        print("File not found. Try again.")
        continue

<<<<<<< HEAD
    if file_type == "html":
        add_html(file_path)
    elif file_type == "docx":
        add_docx(file_path)
    elif file_type == "txt":
        add_txt(file_path)
    else:
        print("Invalid type.")

# Chat Loop
print("\nRAG READY — type exit to quit or 'file' to add a new file")

while True:
    user_input = input("You: ").strip()
    
    if user_input.lower() == 'exit':
        print("RAG: Goodbye!")
        break
    elif user_input.lower() == 'file':  # added incase the user wants to go back and add another file
        while True:
            add_file = input("\nWhat type of file do you want to add?(html/docx/txt/none): ").strip().lower()
            if add_file == "none":
                break

            file_path = input("Enter the path to your file: ").strip().strip('\'"')
=======
    if add_file == "html":
        add_html(file_path)
    elif add_file == "docx":
        add_docx(file_path)
    elif add_file == "txt":
        add_txt(file_path)
    else:
        print("Type html, docx, txt, or none")

# Chat loop
print("\nRAG READY. Type exit to quit or file to add a new file")

while True:
    user_input = input("You: ")
    if user_input.lower() == 'exit':
        print("RAG: Goodbye!")
        break
    elif user_input.lower() == 'file': # added incase the user wants to go back and add another file
        while True:
            add_file = input("\nWhat type of file do you want to add? (html/docx/txt/none): ").strip().lower()
            if add_file == "none":
                break
            file_path = input("Enter the path to your file: ").strip()
            
>>>>>>> 84ff12a2e0009c059180310591bd9e2e38b2253c
            if not os.path.exists(file_path):
                print("File not found. Try again.")
                continue

            if add_file == "html":
                add_html(file_path)
            elif add_file == "docx":
                add_docx(file_path)
            elif add_file == "txt":
                add_txt(file_path)
            else:
<<<<<<< HEAD
                print("Type html, docx, txt, or none")
    else:
        answer = chat(user_input)
        print("RAG:", answer, "\nAnything else needed?")
=======
                print("Type html, docx, txt, or none")     
    else: # any time i typed none, the bot replied immediatly, had to add an else here to fix it
        assistant_rply = chat(user_input)
        print("RAG:", assistant_rply, "\nAnything else needed?")

>>>>>>> 84ff12a2e0009c059180310591bd9e2e38b2253c
